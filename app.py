from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload, MediaIoBaseUpload
import io, os, uuid, base64
from pydantic import BaseModel


app = FastAPI()

# Setup: Create temp directory
UPLOAD_DIR = "temp"
os.makedirs(UPLOAD_DIR, exist_ok=True)


if not os.path.exists("service_account.json"):
    encoded_creds = os.getenv("GOOGLE_CREDS_B64")
    if encoded_creds:
        with open("service_account.json", "wb") as f:
            f.write(base64.b64decode(encoded_creds))
    else:
        raise Exception("GOOGLE_CREDS_B64 env var not set or empty!")

SCOPES = ['https://www.googleapis.com/auth/drive']
creds = service_account.Credentials.from_service_account_file("service_account.json", scopes=SCOPES)
drive_service = build('drive', 'v3', credentials=creds)

FOLDER_ID = "1uSkl46cSRdLvcfXOIjL6RPqEq2vy0CM_"  # your shared Drive folder ID

# Upload DOCX file
def upload_file_to_folder(file_path, file_name, mimetype, folder_id):
    metadata = {'name': file_name, 'parents': [folder_id]}
    media = MediaFileUpload(file_path, mimetype=mimetype)
    uploaded = drive_service.files().create(body=metadata, media_body=media, fields='id').execute()
    drive_service.permissions().create(fileId=uploaded['id'], body={'role': 'reader', 'type': 'anyone'}).execute()
    return f"https://drive.google.com/uc?id={uploaded['id']}&export=download"

# Upload as Google Doc
def upload_as_gdoc(file_path, file_name, folder_id):
    metadata = {'name': file_name, 'parents': [folder_id], 'mimeType': 'application/vnd.google-apps.document'}
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    uploaded = drive_service.files().create(body=metadata, media_body=media, fields='id').execute()
    return uploaded['id']

# Export Google Doc as PDF
def export_gdoc_to_pdf(gdoc_id, pdf_name, folder_id):
    request = drive_service.files().export_media(fileId=gdoc_id, mimeType='application/pdf')
    pdf_data = io.BytesIO()
    downloader = MediaIoBaseDownload(pdf_data, request)
    while True:
        status, done = downloader.next_chunk()
        if done:
            break
    pdf_data.seek(0)
    metadata = {'name': pdf_name, 'parents': [folder_id]}
    media = MediaIoBaseUpload(pdf_data, mimetype='application/pdf', resumable=True)
    uploaded = drive_service.files().create(body=metadata, media_body=media, fields='id').execute()
    drive_service.permissions().create(fileId=uploaded['id'], body={'role': 'reader', 'type': 'anyone'}).execute()
    return f"https://drive.google.com/uc?id={uploaded['id']}&export=download"

# Placeholder replacement with styles


def replace_with_preserved_formatting(doc, replacements, style_overrides):
    def replace_in_runs(runs):
        for run in runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    if key in style_overrides:
                        run.font.name = style_overrides[key].get('name', run.font.name)
                        run.font.size = style_overrides[key].get('size', run.font.size)

    for para in doc.paragraphs:
        replace_in_runs(para.runs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para.runs)
# Main function
def generate_and_upload(details):
    template_path = "final_1.docx"
    output_docx = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}.docx")

    loi_date = datetime.strptime(details['LOI due date'], "%m/%d/%Y")
    replacements = {
        'name_1': details['Closer 1 name'],
        'name_2': details['Closer 2 name'],
        'name1_2': details['Closer 1 name'],
        'name2_2': details['Closer 2 name'],
        '[Closer 1 Name_3]': details['Closer 1 name'],
        '[Closer 2 Name_3]': details['Closer 2 name'],
        '[Closer 1 Title]': details['Closer 1 title'],
        '[Closer 2 Title]': details['Closer 2 title'],
        'due_date': details['LOI due date'],
        '[Company’s Name]': details['Company name'],
        'email_1': details['Closer 1 email'],
        'email_2': details['Closer 2 email'],
        '[Closer 1 Number]': details['Closer 1 number'],
        '[Closer 2 Number]': details['Closer 2 number'],
        '[Today’s Date]': datetime.now().strftime("%m/%d/%Y"),
        '[LOI Due Date +7]': (loi_date + timedelta(days=7)).strftime("%m/%d/%Y"),
        '[LOI Due Date +67]': (loi_date + timedelta(days=67)).strftime("%m/%d/%Y"),
    }

    styles = {
        'name_1': {'name': 'Arial', 'size': Pt(12)},
        'name_2': {'name': 'Arial', 'size': Pt(12)},
        'due_date': {'name': 'Arial', 'size': Pt(12), 'bold': True},
        '[Company’s Name]': {'name': 'Arial', 'size': Pt(11)},
        '[Today’s Date]': {'name': 'Arial', 'size': Pt(11)},
        'email_1': {'name': 'Arial', 'size': Pt(12)},
        'email_2': {'name': 'Arial', 'size': Pt(12)},
        'name1_2': {'name': 'Palatino Linotype', 'size': Pt(30), 'italic': True},
        'name2_2': {'name': 'Palatino Linotype', 'size': Pt(30), 'italic': True},
        '[Closer 1 Name_3]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[Closer 2 Name_3]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[Closer 1 Title]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[Closer 2 Title]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[Closer 1 Number]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[Closer 2 Number]': {'name': 'Times New Roman', 'size': Pt(10)},
        '[LOI Due Date +7]': {'name': 'Arial', 'size': Pt(11)},
        '[LOI Due Date +67]': {'name': 'Arial', 'size': Pt(11)},
    }

    doc = Document(template_path)
    replace_with_preserved_formatting(doc, replacements, styles)
    doc.save(output_docx)

    docx_link = upload_file_to_folder(output_docx, "Generated_LOI.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", FOLDER_ID)
    gdoc_id = upload_as_gdoc(output_docx, 'Generated_LOI_GDoc', FOLDER_ID)
    pdf_link = export_gdoc_to_pdf(gdoc_id, 'Generated_LOI.pdf', FOLDER_ID)

    return {
        "docx": docx_link,
        "pdf": pdf_link
    }

class LOIDetails(BaseModel):
    Closer_1_name: str
    Closer_2_name: str
    Closer_1_title: str
    Closer_2_title: str
    Closer_1_email: str
    Closer_2_email: str
    Closer_1_number: str
    Closer_2_number: str
    Company_name: str
    LOI_due_date: str  # format: YYYY-MM-DD

@app.post("/generate/")
async def generate_loi(data: LOIDetails):
    # Convert to dict and restore original keys
    details = {
        "Closer 1 name": data.Closer_1_name,
        "Closer 2 name": data.Closer_2_name,
        "Closer 1 title": data.Closer_1_title,
        "Closer 2 title": data.Closer_2_title,
        "Closer 1 email": data.Closer_1_email,
        "Closer 2 email": data.Closer_2_email,
        "Closer 1 number": data.Closer_1_number,
        "Closer 2 number": data.Closer_2_number,
        "Company name": data.Company_name,
        "LOI due date": data.LOI_due_date
    }

    result = generate_and_upload(details)
    return JSONResponse(content=result)
